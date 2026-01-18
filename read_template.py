from docx import Document

doc = Document('Template for Reference.docx')

print("=== TEMPLATE STRUCTURE ===\n")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else "Normal"
        print(f"[{style}] {para.text}")
        
print("\n\n=== TABLE CONTENT ===\n")
for table in doc.tables:
    for row in table.rows:
        print(" | ".join([cell.text.strip() for cell in row.cells]))
    print()
